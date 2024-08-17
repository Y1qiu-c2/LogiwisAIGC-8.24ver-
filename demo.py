import os
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer


def read_docx(file_path):
    """读取.docx文件的内容，包括标题和全文"""
    doc = Document(file_path)

    # 假设标题是文档的第一个段落（如果有）
    title = doc.paragraphs[0].text if doc.paragraphs else ""

    # 提取正文内容
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    content = "\n".join(full_text)
    return title, content


def extract_keywords(query):
    """提取查询中的关键词"""
    vectorizer = CountVectorizer(stop_words='english')
    X = vectorizer.fit_transform([query])
    return vectorizer.get_feature_names_out()


def retrieve_related_documents(query, folder_path):
    """根据用户查询在指定文件夹中检索相关的.docx文件，考虑标题和内容"""
    related_files = []

    # 提取用户查询中的关键词
    keywords = extract_keywords(query)

    # 遍历文件夹中的所有.docx文件
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            title, content = read_docx(file_path)

            # 检查任何一个关键词是否出现在标题或正文中
            if any(keyword.lower() in title.lower() or keyword.lower() in content.lower() for keyword in keywords):
                related_files.append(file_path)

    return related_files


# 示例用法
folder_path = "/Users/qiuyipeng/Desktop/RAGDATA"
query = input("关键词：")
related_documents = retrieve_related_documents(query, folder_path)

# 输出检索到的相关文件
for file_path in related_documents:
    print(f"File: {file_path}")